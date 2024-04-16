# ==== Google Calendar API reqs =====
from __future__ import print_function

import datetime, json, traceback, sys, re, pytz, calendar
from datetime import timedelta
import os.path
from .models import Calendar, Shifts
from django.utils import timezone
from django.db.models import Q
from django.shortcuts import render, redirect, HttpResponse
from django.http import JsonResponse
from docx import Document
import csv, json
from dateutil.parser import parse

# from google.auth.transport.requests import Request
# from google.oauth2.credentials import Credentials
# from google_auth_oauthlib.flow import InstalledAppFlow
# from googleapiclient.discovery import build
# from googleapiclient.errors import HttpError
# ================================================================


# ======== NOTE: Need to update this so Office Manager can set Hospital Timezone in Admin Settings =========
TIMEZONE = pytz.timezone('America/Los_Angeles')

month_variables = ["01", "02", "03", "04", "05", "06", "07", "08", "09", "10", "11", "12"]

# wordDoc = Document(filedialog.askopenfilename( filetypes = ( (".docx .doc files", "*.docx *.doc"),("All files", "*.*") ) ))

month_abbrev = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]

month_list = [
    "january", "jan",
    "february", "feb",
    "march", "mar",
    "april", "apr",
    "may",
    "june", "jun",
    "july", "jul",
    "august", "aug",
    "september", "sept",
    "october", "oct",
    "november", "nov",
    "december", "dec"
  ]

day_list = [
    "monday", "mon",
    "tuesday", "tues",
    "wednesday", "wed",
    "thursday", "thurs",
    "friday", "fri",
    "saturday", "sat",
    "sunday", "sun"
  ]

FORM_FIELD_TYPES = {
  'shift': 'select',
  'shift_type': 'select',
  'user': 'select',
  'shift_date': 'date',
  'shift_name': 'input',
  'shift_label': 'input',
  'start_time': 'time',
  'end_time': 'time',
  'shift_color': 'color',
  'permissions': 'multi-select',
}

FORM_FIELD_LABELS = {
  'shift': 'Shift',
  'shift_type': 'Shift Type',
  'user': 'User',
  'shift_date': 'Shift Date',
}

REQUIRED_FIELDS = [
  'shift',
  'shift_name',
  'shift_type',
  'user',
  'shift_date',
]

REMOVE_FIELDS = [
  'created_at',
  'updated_at',
]

def fix_timezone(dt):
    # Assume dt is a naive datetime object in UTC
    dt = timezone.make_aware(dt, timezone=pytz.UTC)
    # Convert to the 'America/Los_Angeles' timezone
    dt = dt.astimezone(TIMEZONE)
    return dt

def trace_error(e, isForm=False):
    exc_type, exc_value, exc_traceback = sys.exc_info()
    filename, line_number, func_name, text = traceback.extract_tb(exc_traceback)[0]
    print(f"An error occurred in file {filename} on line {line_number} in {func_name}(): {text}")
    print("Error: ", e)
    if isForm:
        return JsonResponse({'message':'Form is invalid'}, status=500)
    return JsonResponse({'message':'Something went wrong'}, status=500)

def convert_schedule(schedule, user, month, year):
  user_month = month
  user_year = year
  wordDoc = Document(schedule)
  while not user_month in month_variables:
    # user_month = simpledialog.askstring(title="Month", prompt="Please enter the 2-digit month")
    user_month = "08"

  while not int(user_year) > 2021:
    # user_year = simpledialog.askstring(title="Year", prompt="Please enter the 4-digit year")
    user_year = "2022"
  # wordDoc = Document('SEP21_schedule.docx')
  # wordDoc = Document('Aug 2022(CORRECT).docx')
  month = user_month
  year = "2022"
  shifts = []
  for table in wordDoc.tables:
      date = []
      j = 0
      for row in table.rows:
        row_text = ''
        i = 0
        k = 0
        shift = ''
        time = ''
        for cell in row.cells:
          row_text = row_text + cell.text + ","
          if cell.text.lower() in day_list:
            j = 0
            # print('reading cells')
          
          if j % 6 == 1:
              date.append(cell.text)
          else:
            if cell.text in month_list:
              month = cell.text
            if i % 2 == 0:
              shift = cell.text
              if j % 6 == 2:
                time = "07:00"
              if j % 6 == 3:
                time = "10:00"
              if j % 6 == 4:
                time = "14:00"
              if j % 6 == 5:
                time = "18:00"
            # if "MA" in cell.text:
            # print(user)
            if user in cell.text:
              shift_start = datetime.datetime(
                int(user_year),
                int(user_month),
                int(date[i]),
                int(time[:2]),
                00,
                00
              )
              # print(shift_start)
              shifts.append({
                "shift": shift, 
                "date": user_year + "-" + user_month + "-" + date[i], 
                "time": time, 
                "shift_start": str(shift_start),
                "shift_end": str(shift_start + timedelta(hours=12)),
              })
                
          i += 1

        # print(row_text, "row", j)
        # print("date row", date)
        j += 1
        i = 0
        k = 0
        if j % 6 == 0: 
          date = []

  header = ['Subject', 'Start date', 'Start time']
  # with open(f'schedule_{user_month}-{user_year}.csv', 'w', encoding='UTF8', newline='\n') as f:
  #   writer = csv.writer(f)
  #   writer.writerow(header)
  #   for shift in shifts:
  #     writer.writerow(shift)

  # f.close()
  # start_time = datetime.datetime(2023,6,21,3,45,00)
  # end_time = start_time + timedelta(hours=12)
  # print("start: ", start_time, "end: ", end_time)
  # print(shifts)
  
  events = add_shifts(shifts)
  json_shifts = json.dumps(events)
  return json_shifts

def load_schedule(schedule, month, year, users = None, shift_names = None):
  # print(f'year start: {year}')
  # print(f'users: {users} \n shift_names: {shift_names} \n')
  # print('JM: ', users['JM'])
  # print('Swing 1 start: ', shift_names[1]['start'])
  user_list = set([])
  user_month = month
  user_year = year
  first_day = datetime.datetime.strptime(f'{year}-{month}-01', '%Y-%m-%d').date()
  _, last_day_num = calendar.monthrange(first_day.year, first_day.month)
  last_day = datetime.datetime(first_day.year, first_day.month, last_day_num).date()
  # print(f'first day: {first_day} \n last day: {last_day}')
  wordDoc = Document(schedule)
  # user_tz = pytz.timezone('America/Los_Angeles')
  while not user_month in month_variables:
    # user_month = simpledialog.askstring(title="Month", prompt="Please enter the 2-digit month")
    user_month = "08"
  while not int(user_year) > 2021:
    user_year = "2022"
  month = user_month
  # year = "2022"
  shifts = []
  remove_shifts = set([])
  # ================= Pull Current Shifts =================
  current_shifts = list(Shifts.objects.filter(shift_start__year=year, shift_start__month=month).values())
  # shift_times = {2: "07:00", 3: "10:00", 4: "14:00", 5: "18:00"}
  # print(f'shift_names: {shift_names}')
  shift_times = {2: shift_names[0], 3: shift_names[1], 4: shift_names[2], 5: shift_names[3]}
  # print(f'shift_times: {shift_times}')
  shift_name_ids = {2: shift_names[0]['id'], 3: shift_names[1]['id'], 4: shift_names[2]['id'], 5: shift_names[3]['id']}
  for table in wordDoc.tables:
      date = []
      j = 0
      for row in table.rows:
        row_text = ''
        i = 0
        k = 0
        shift = ''
        time = ''        
        for cell in row.cells:
          cell.text = cell.text.replace(' ', '')
          row_text = row_text + cell.text + ","
          if cell.text in month_list:
            month = cell.text
          if cell.text.lower() in day_list:
            j = 0
            # print('reading cells')
          else:
            if j % 6 == 1:
                date.append(cell.text)
                # print(f'date {date}')
            # else:
            #   if i % 2 == 0:
            else: 
              if "/" in cell.text:
                user_cell = cell.text.split("/")
              else:
                user_cell = [cell.text]
              for user in user_cell:
                if user != "" and user in users:
                  try:
                    user_list.add(users[user])
                    start = datetime.datetime.strptime(shift_times[j % 6]['start'], '%H:%M').time()
                    end = datetime.datetime.strptime(shift_times[j % 6]['end'], '%H:%M').time()
                    # print(start, end)
                    shift_start = convert_to_shift_datetime(f'{user_year}-{user_month}-{date[i]}', start)
                    if start > end:
                      shift_end = convert_to_shift_datetime(f'{user_year}-{user_month}-{date[i]}', end) + timedelta(days=1)
                    else:
                      shift_end = convert_to_shift_datetime(f'{user_year}-{user_month}-{date[i]}', end)
                    # print(f'shift_start: {shift_start} \n shift_end: {shift_end}')
                    # user_date = start.replace(tzinfo=TIMEZONE)
                    # filter_date = user_date.astimezone(pytz.timezone('UTC')).date()
                    filter_date = shift_start.date()
                    print(f'filter_date: {filter_date}')
                    # ================= individually update pulled shifts and create new shifts =================
                    # filtered_shift = current_shifts.filter(user_id=users[user], shift_start__date=filter_date).values()
                    filtered_shift = list(filter(lambda x: x['user_id'] == users[user] and x['shift_start'].date() == filter_date, current_shifts))
                    shift_type = 1
                    if len(filtered_shift) > 0 :
                      shift_type = filtered_shift[0]['shift_type_id']
                      for shift in filtered_shift:
                        remove_shifts.add(shift['id'])
                      # new_shift = {}
                      # print(f'filtered_shift: {filtered_shift}')
                      # new_shift['shift_start'] = shift_start,
                      # new_shift['shift_end'] = shift_end,
                      # new_shift['shift_name_id'] = shift_name_ids[j % 6],
                      # new_shift['user_id'] = users[user],
                      # new_shift['shift_type_id'] = 1,
                      # print("updated shift: ", new_shift)
                    else:
                      print(f'no filtered shift found for {users[user]} on {filter_date}')
                    
                    # =============== Used to individually create/update shifts ===============
                    # shift, created = Shifts.objects.update_or_create(
                    #   user_id = users[user],
                    #   shift_start__date=filter_date,
                    #   defaults = {
                    #     'shift_start': shift_start,
                    #     'shift_end': shift_end,
                    #     'shift_name_id': shift_name_ids[j % 6],
                    #     'user_id': users[user],
                    #     'shift_type_id': 1,
                    #   }
                    # )
                    # ================= Used to bulk create shifts =================
                    shifts.append({
                      'shift_start': shift_start,
                      'shift_end': shift_end,
                      'shift_name_id': shift_name_ids[j % 6],
                      'user_id': users[user],
                      'shift_type_id': shift_type,
                    })               

                  except ValueError:
                    print(f"Invalid date: {user_year}-{user_month}-{date[i]}", ValueError)
          i += 1

        # print(row_text, "row", j)
        # print("date row", date)
        j += 1
        i = 0
        k = 0
        if j % 6 == 0: 
          date = []
  user_list = list(user_list)
  remove_shifts = list(remove_shifts)
  # current_shifts = list(filter(lambda x: x['id'] not in remove_shifts, current_shifts))
  # print(f'users from upload: {user_list}')
  # ================== Used to bulk remove then bulk create shifts ==================
  Shifts.objects.filter(
    Q(id__in=remove_shifts),
  ).delete()
  Shifts.objects.bulk_create(
    Shifts(**shift) for shift in shifts
  )
  # ================== Used to bulk create shifts ==================
  # Shifts.objects.filter(
  #   Q(user_id__in=user_list),
  #   Q(shift_start__date__gte=first_day), 
  #   Q(shift_end__date__lte=last_day),
  # ).delete()
  # Shifts.objects.bulk_create(
  #   Shifts(**shift) for shift in shifts
  # )
  return "Success!"

def load_schedule_old(schedule, month, year):
  print(f'year start: {year}')
  user_month = month
  user_year = year
  wordDoc = Document(schedule)
  # user_tz = pytz.timezone('America/Los_Angeles')
  while not user_month in month_variables:
    # user_month = simpledialog.askstring(title="Month", prompt="Please enter the 2-digit month")
    user_month = "08"

  while not int(user_year) > 2021:
    user_year = "2022"
  month = user_month
  # year = "2022"
  shifts = []
  shift_times = {2: "07:00", 3: "10:00", 4: "14:00", 5: "18:00"}
  for table in wordDoc.tables:
      date = []
      j = 0
      for row in table.rows:
        row_text = ''
        i = 0
        k = 0
        shift = ''
        time = ''
        for cell in row.cells:
          row_text = row_text + cell.text + ","
          if cell.text in month_list:
            month = cell.text
          if cell.text.lower() in day_list:
            j = 0
            # print('reading cells')
          else:
            if j % 6 == 1:
                date.append(cell.text)
                # print(f'date {date}')
            # else:
            #   if i % 2 == 0:
            elif i % 2 == 0:
                shift = cell.text
                # if j % 6 == 2:
                #   time = "07:00"
                # if j % 6 == 3:
                #   time = "10:00"
                # if j % 6 == 4:
                #   time = "14:00"
                # if j % 6 == 5:
                #   time = "18:00"
              # print(user)
            elif cell.text != "":
              cell_user = cell.text
              try:
                shift_start = datetime.datetime(
                  int(user_year),
                  int(user_month),
                  int(date[i]),
                  int(shift_times[j % 6][:2]),
                  00,
                  00
                )
                # shift_start = timezone.make_aware(shift_start, timezone=TIMEZONE)
                # print(shift_start)
                # print(f'user: {cell_user}')
                shifts.append({
                  "user": cell_user,
                  "shift": shift, 
                  "date": user_year + "-" + user_month + "-" + date[i], 
                  "time": shift_times[j % 6], 
                  "month": month,
                  "year": user_year,
                  "shift_start": str(shift_start),
                  "shift_end": str(shift_start + timedelta(hours=12)),
                })
              except ValueError:
                print(f"Invalid date: {user_year}-{user_month}-{date[i]}")
          i += 1

        # print(row_text, "row", j)
        # print("date row", date)
        j += 1
        i = 0
        k = 0
        if j % 6 == 0: 
          date = []

  # header = ['Subject', 'Start date', 'Start time']
  # with open(f'schedule_{user_month}-{user_year}.csv', 'w', encoding='UTF8', newline='\n') as f:
  #   writer = csv.writer(f)
  #   writer.writerow(header)
  #   for shift in shifts:
  #     writer.writerow(shift)

  # f.close()
  # start_time = datetime.datetime(2023,6,21,3,45,00)
  # end_time = start_time + timedelta(hours=12)
  # print("start: ", start_time, "end: ", end_time)
  # print(shifts)
  load_database(shifts, user_month, user_year)
  # events = add_shifts(shifts)
  # json_shifts = json.dumps(events)
  return "Success!"

def load_database(shifts, month, year):
  try:
    print(month, year)
    Calendar.objects.filter(month=month, year=year).delete()
    i = 0
    for shift in shifts:
      calendar_shift = Calendar.objects.create(
        user_initials = shift["user"].strip(),
        start = shift["shift_start"],
        end = shift["shift_end"],
        month = month,
        year = year,
      )
    return
  except Exception as e:
    print('An error occurred: %s' % e)  
    return trace_error(e, True)

SCOPES = ['https://www.googleapis.com/auth/calendar']

def get_users(schedule):
  wordDoc = Document(schedule) # Identify schedule as a Word Doc
  user_list = []
  for table in wordDoc.tables:
    for row in table.rows:
      row_text = ''
      for cell in row.cells:
        row_text = row_text + cell.text + ","
        if cell.text.isalpha() and len(cell.text) == 2:
          # print(f'cell: {cell.text}')
          if cell.text not in user_list:          
            user_list.append(cell.text)
    # print(user_list)
    # return json.dumps(user_list)
    return user_list


def add_shifts(shifts):
  print('running add_events')
  creds = None

  try:
      # service = build('calendar', 'v3', credentials=creds)
      i = 0
      events = {}
      for shift in shifts:
        # print("first half: ", shift["shift_start"][:10])
        # print("second half: ", shift["shift_start"][11:])
        # events = {
        events[f'shift_{i}'] = {
          "summary": shift["user"], # shift["shift"] + "-" + shift["user"],
          "location": "AMCS",
          "description": shift["shift"],
          "start": {
            "dateTime": shift["shift_start"][:10] + "T" + shift["shift_start"][11:] + "-07:00",
            "timeZone": "America/Los_Angeles"
          },
          "end": {
            "dateTime": shift["shift_end"][:10] + "T" + shift["shift_end"][11:] + "-07:00",
            "timeZone": "America/Los_Angeles"
          },
        }
        i += 1
        # print(event["start"])
        # events = service.events().insert(calendarId='primary', body=event).execute()
      # print("new events: ", events)
      return events
  except Exception as e:
    print('An error occurred: %s' % e)  
    return trace_error(e, True)
  
SEARCH_LIST = [
  'permission',
]

def find_locked(array1, array2 = SEARCH_LIST):
  # Convert both arrays to sets
  set1 = set(array1)
  set2 = set(array2)
  # Find the intersection of the two sets
  common_elements = set1 & set2

  # If there are at least two common elements, trigger the function
  if len(common_elements) >= 2:
    return True
  return False

def set_form_fields(form, model=''):
  new_form = {}
  print("===============>    ", list(form.fields.keys()))
  field_names = list(form.fields.keys())
  for field in field_names:
    if not field in REMOVE_FIELDS:
      new_form[field] = {      
        'label': FORM_FIELD_LABELS[field] if field in FORM_FIELD_LABELS else convert_label(field),
        'type': FORM_FIELD_TYPES[field] if field in FORM_FIELD_TYPES else 'input',
        'value': None,
        'required': field in REQUIRED_FIELDS if field in REQUIRED_FIELDS else False,
      }    
  if find_locked(field_names):
    new_form['permission']['type'] = 'locked'
  if model == 'ShiftType':
    new_form['shift_type']['type'] = 'input'
  return new_form

def convert_label(name):
    name_with_spaces = re.sub('_', ' ', name)
    capitalized_name = name_with_spaces.title()
    return capitalized_name

def convert_to_shift_datetime(date, time):
  # print(date, time)
  shift_date = parse(date).date()
  shift_datetime = datetime.datetime.combine(shift_date, time)
  # shift_datetime = fix_timezone(shift_datetime)
  return shift_datetime